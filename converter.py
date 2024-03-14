###### CONVERTER
###### PROGRAM WRITTEN BY - Adam Reynoldson reynoldson2002@gmail.com FOR Darren Reid via UpWork 2023
######
###### Code for reading/saving the 'Discovery Responses'
######


### FUNCTIONS
########################################################################################################

from functions import *
import PyPDF2 as pdf
import re,time,random
import copy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.parser import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from pdfminer.high_level import extract_text
import fitz
from CTkMessagebox import CTkMessagebox
from boxdetect.pipelines import get_checkboxes
from boxdetect import config
from boxdetect.pipelines import get_boxes
import cv2
import concurrent.futures


### CONSTANTS
########################################################################################################

FORM_VALUES = ["(1)","(2)",'1.1', '2.1', '2.2', '2.3', '2.4', '2.5', '2.6', '2.7', '2.8', '2.9', '2.10', '2.11', '2.12', '2.13', '3.1', '3.2', '3.3', '3.4', '3.5', '3.6', '3.7', '4.1', '4.2', '6.1', '6.2', '6.3', '6.4', '6.5', '6.6', '6.7', '7.1', '7.2', '7.3', '8.1', '8.2', '8.3', '8.4', '8.5', '8.6', '8.7', '8.8', '9.1', '9.2', '10.1', '10.2', '10.3', '11.1', '11.2', '12.1', '12.2', '12.3', '12.4', '12.5', '12.6', '12.7', '13.1', '13.2', '14.1', '14.2', '15.1', '16.1', '16.2', '16.3', '16.4', '16.5', '16.6', '16.7', '16.8', '16.9','16.10', '17.1', '20.1', '20.2', '20.3', '20.4', '20.5', '20.6', '20.7', '20.8', '20.9','20.10', '20.11', '50.1', '50.2', '50.3', '50.4', '50.5', '50.6']


### FUNCTIONS
########################################################################################################

# Adds paragraph after current
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")

    paragraph._p.addnext(new_p)
    
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

# Deletes a paragraph
def delete_paragraph(paragraph):
    p = paragraph._element
    if p!=None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

### READ THE PDF
########################################################################################################

#Reads a pdf use pdfreader
def readPDF(file):
    reader = pdf.PdfReader(file)# Create reader object
    #print(reader.get_fields())
    text=[]# Get each page
    for i in reader.pages:
        text.append(i.extract_text())
    text = "".join(list(text))
    return text

#Reads a pdf using fitz
def readPDF3(file):
    doc = fitz.Document(file)
    text = ""
    for page in doc:
        """
        print(page.rect.width)
        crop_rect = fitz.Rect(300,0,page.rect.width,page.rect.height)
        page.set_mediabox(crop_rect)
        page.set_cropbox(crop_rect)
        page.set_trimbox(crop_rect)
        print(page)
        print(page.mediabox)
        """
        #REMOVE VERTICAL TEXT!!
        for block in page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]:
            for line in block["lines"]:
                wdir = line["dir"]    # writing direction = (cosine, sine)
                if wdir[0] == 0:  # either 90° or 270°
                    page.add_redact_annot(line["bbox"])
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)  # remove text, but no image
        

        text+=page.get_text()
        #print(page.rect.width)
    return text

def generate_unique_key():
    # Get current timestamp
    timestamp = int(time.time() * 1000)  # Multiply by 1000 to get milliseconds
    # Generate a random number
    random_number = random.randint(0, 1000000)
    # Combine timestamp and random number to create a unique key
    unique_key = f"{timestamp}-{random_number}"
    return unique_key


### READ FROGs FORM
########################################################################################################
"""
def readForm(file):# file, (w,h,ratio,dilation)
    cfg = config.PipelinesConfig()
    # important to adjust these values to match the size of boxes on your image
    cfg.width_range = (9,30)
    cfg.height_range = (9,30)
    # the more scaling factors the more accurate the results but also it takes more time to processing
    # too small scaling factor may cause false positives
    # too big scaling factor will take a lot of processing time
    cfg.scaling_factors = [10]
    # w/h ratio range for boxes/rectangles filtering
    cfg.wh_ratio_range = (0.8, 2.5)
    # group_size_range starting from 2 will skip all the groups
    # with a single box detected inside (like checkboxes)
    cfg.group_size_range = (1, 1)
    # num of iterations when running dilation tranformation (to engance the image)
    cfg.dilation_iterations = 0

    results=[]
    vals=[]
    f = fitz.open(file)
    page_count = 0#Used to add extra results when scanning fails
    sizes=[]
    for page in f:
        img = page.get_pixmap()
        img.save(os.path.join(os.path.dirname(__file__),"assets/out.png"))
        #SPLIT INTO COLUMNS
        img = cv2.imread(os.path.join(os.path.dirname(__file__),"assets/out.png"))
        height = img.shape[0]
        width = img.shape[1]

        # Cut the image in half
        width_cutoff = int(width // 2.2)
        s1 = img[:, :width_cutoff]
        s2 = img[:, width_cutoff:]
        s1 = img[:, :100]
        s2 = img[:, width_cutoff:width_cutoff+150]

        cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/s1.png"), s1)
        cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/s2.png"), s2)

        checkboxes1 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/s1.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)
        checkboxes2 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/s2.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)

        for checkboxes in [checkboxes1,checkboxes2]:
            #GET TRUE OF FALSE FOR CHECKBOXES ON THE PAGE

            for check in checkboxes:
                print("Checkbox bounding rectangle (x,y,width,height): ", check[0])
                sizes.append(check[0])

                #Get array
                array = check[2]
                total=0
                for row in array[2:-4,2:-4]:
                    total+=sum(row)/len(row)
                total = total/len(array[2:-4])

                vals.append(total)
                
                if total>=20:
                    results.append(True)
                else:
                    results.append(False)
        if page_count==0 and len(results)==0:
            results.append(False)
    output=[]

    #IF NOT ENOUGH, ENHANCE AND DO AGAIN WITH EXACT SHAPE!
    if len(results)<90:
        print(len(results))
        print("NOT ENOUGH DETECTED: TRYING AGAIN")
        avg_width = int(sum(item[2] for item in sizes[5:20])/len(sizes[5:20]))
        avg_height = int(sum(item[3] for item in sizes[5:20])/len(sizes[5:20]))
        ratio = avg_width/avg_height
        print(ratio)
        print(avg_width,avg_height)
        cfg.width_range = (avg_width-5,avg_width+5)
        cfg.height_range = (avg_height-5,avg_height+5)
        cfg.wh_ratio_range = (ratio-0.4, ratio+0.4)
        cfg.dilation_iterations = 1
        cfg.scaling_factors=[10]

        results=[]
        vals=[]
        f = fitz.open(file)
        page_count = 0#Used to add extra results when scanning fails
        sizes=[]
        for page in f:
            img = page.get_pixmap()
            img.save(os.path.join(os.path.dirname(__file__),"assets/out.png"))
            #SPLIT INTO COLUMNS
            img = cv2.imread(os.path.join(os.path.dirname(__file__),"assets/out.png"))
            height = img.shape[0]
            width = img.shape[1]

            # Cut the image in half
            width_cutoff = int(width // 2.5)
            s1 = img[:, :200]
            s2 = img[:, width_cutoff:width_cutoff+200]

            cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/s1.png"), s1)
            cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/s2.png"), s2)

            checkboxes1 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/s1.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)
            checkboxes2 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/s2.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)

            for checkboxes in [checkboxes1,checkboxes2]:
                #GET TRUE OF FALSE FOR CHECKBOXES ON THE PAGE

                for check in checkboxes:
                    sizes.append(check[0])

                    #Get array
                    array = check[2]
                    total=0
                    for row in array[2:-4,2:-4]:
                        total+=sum(row)/len(row)
                    total = total/len(array[2:-4])

                    vals.append(total)
                    
                    if total>=20:
                        results.append(True)
                    else:
                        results.append(False)
            if page_count==0 and len(results)==0:
                results.append(False)
    
    print(len(results))
    results = results[:90]
    print(len(FORM_VALUES))


readForm(str("C:/Users/Adam/Desktop/Freelancing/Darren Reid PDF/Darren Reid PDF/FROGs for Adam/SERVED - Def FROGs (Set 1) to Plf.pdf"))
"""
### READ FROGs FORM THREADED!!!!
########################################################################################################

def formThread(page,cfg):
    #SAVE PDF PAGE AS IMAGE
    checkbox_sizes=[]
    tag = str(generate_unique_key())
    local_results=[]
    img = page.get_pixmap()
    img.save(os.path.join(os.path.dirname(__file__),"assets/"+tag+"out.png"))

    #CUT IMAGE IN HALF
    img = cv2.imread(os.path.join(os.path.dirname(__file__),"assets/"+tag+"out.png"))
    height = img.shape[0]
    width = img.shape[1]
    width_cutoff = int(width // 2.5)
    s1 = img[:, :200]
    s2 = img[:, width_cutoff:width_cutoff+200]
    cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/"+tag+"s1.png"), s1)
    cv2.imwrite(os.path.join(os.path.dirname(__file__),"assets/"+tag+"s2.png"), s2)

    #GET CHECKBOXES
    checkboxes1 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/"+tag+"s1.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)
    checkboxes2 = get_checkboxes(os.path.join(os.path.dirname(__file__),"assets/"+tag+"s2.png"), cfg=cfg, px_threshold=0.1, plot=False, verbose=False)

    #FILTERING
    for checkboxes in [checkboxes1,checkboxes2]:
        #GET TRUE OF FALSE FOR CHECKBOXES ON THE PAGE
        for check in checkboxes:
            #Get array
            checkbox_sizes.append(check[0])
            array = check[2]
            #Get an average pixel value
            total=0
            for row in array[2:-4,2:-4]:
                total+=sum(row)/len(row)
            total = total/len(array[2:-4])

            #If above a threshold then True
            if total>=20:
                local_results.append(True)
            else:

                local_results.append(False)
    
    #DELETE SAVED IMAGES
    files = [os.path.join(os.path.dirname(__file__),"assets/"+tag+"out.png"),
             os.path.join(os.path.dirname(__file__),"assets/"+tag+"s1.png"),
             os.path.join(os.path.dirname(__file__),"assets/"+tag+"s2.png")]
    for file in files:
        if os.path.exists(file):
            os.remove(file)

    return local_results,checkbox_sizes


def readFormThreaded(file):
    try:
        #1. SETTING UP BOXDETECT PIPELINE
        cfg = config.PipelinesConfig()
        # important to adjust these values to match the size of boxes on your image
        cfg.width_range = (9,30)
        cfg.height_range = (9,30)
        # the more scaling factors the more accurate the results but also it takes more time to processing
        # too small scaling factor may cause false positives
        # too big scaling factor will take a lot of processing time
        cfg.scaling_factors = [1]
        # w/h ratio range for boxes/rectangles filtering
        cfg.wh_ratio_range = (0.8, 2.5)
        # group_size_range starting from 2 will skip all the groups
        # with a single box detected inside (like checkboxes)
        cfg.group_size_range = (1, 1)
        # num of iterations when running dilation tranformation (to engance the image)
        cfg.dilation_iterations = 0

        #2. USE THREADING TO FIND CHECKBOXES
        f = fitz.open(file)
        futures=[]
        results=[]
        checkbox_sizes=[]
        with concurrent.futures.ThreadPoolExecutor() as executor:
            for page in f:
                future = executor.submit(formThread, page,cfg)
                futures.append(future)
            c=0
            for future in futures:
                result = future.result()[0]
                sizes_temp = future.result()[1]
                results.extend(result) 
                checkbox_sizes.extend(sizes_temp)
                if c==0 and results==[]:
                    results.append(False) 
                c+=1


        #3. IF NOT ENOUGH DETECTED, ENHANCE AND DO AGAIN WITH EXACT SHAPE!
        
        if len(results)<90:
            print("NOT ENOUGH DETECTED: TRYING AGAIN")
            avg_width = int(sum(item[2] for item in checkbox_sizes[5:20])/len(checkbox_sizes[5:20]))
            avg_height = int(sum(item[3] for item in checkbox_sizes[5:20])/len(checkbox_sizes[5:20]))
            ratio = avg_width/avg_height
            print(ratio)
            print(avg_width,avg_height)
            cfg.scaling_factors = [10]

            #Repeat attempt but with fine tuned parameters
            futures=[]
            results=[]
            checkbox_sizes=[]
            with concurrent.futures.ThreadPoolExecutor() as executor:
                for page in f:
                    future = executor.submit(formThread, page,cfg)
                    futures.append(future)
                c=0
                for future in futures:
                    result = future.result()[0]
                    sizes_temp = future.result()[1]
                    results.extend(result) 
                    checkbox_sizes.extend(sizes_temp)
                    if c==0 and results==[]:
                        results.append(False) 
                    c+=1
        

        #4. ONLY TAKE THE FIRST 90 VALUES!!!!
        output=[]
        print(len(results))
        results = results[:90]
        print(len(FORM_VALUES))

        if len(results)!=len(FORM_VALUES):#IF THE NUMBER OF DETECTED BOXES IS NOT CORRECT
            msg = CTkMessagebox(title="Could not load FROG", message="An incorrect number of FROG's was detected. Defaulted to using all discovery requests.",
                                icon="warning", option_1="Okay",corner_radius=0)
            output = FORM_VALUES.copy()
            return output
        
        #5. GET THE FORM VALUES!
        for i in range(len(FORM_VALUES)):
            if results[i]==True:
                output.append(FORM_VALUES[i])

    except Exception as e:
        print(e)
        msg = CTkMessagebox(title="Could not load FROG", message="The FROG file could not be read correctly. Defaulted to using all discovery requests.",
                            icon="warning", option_1="Okay",corner_radius=0)
        output = FORM_VALUES.copy()
    print("In Form")
    print(output)
    #6. RETURN THE OUTPUT VALUE!
    return output






### FILTER THROUGH THE PDF
########################################################################################################
# Get requests and details from a pdf text string
def filterPDF(data):
    #Search Terms
    terms=["REQUESTNO.","INTERROGATORYNO.","ANDNO.","ENTSNO.","ONSNO.","TIONNO.","REQUESTFORADMISSION"]
    #Found Requests
    reqs=[]
    #Found Keys
    keys=[]
    #Current request
    req=""
    #Adding to the current request? (spans multiple lines)
    adding=False
    term_used=False
    #If program must stop
    hard_stop=False
    start_at_one=True
    #Split of ALL the data in the PDF
    split = data.split("\n")
    #List of document details
    case_number=""
    county=""
    plaintiff=""
    defendant=""
    document=""
    propounding_party=""
    responding_party=""

    ########################################## MAIN SEARCH LOOP

    for i in range(len(split)):
        check=2
        ##################################### 1. GET DOCUMENT DETAILS!
        try:
            if adding==False:#THIS KEEPS THE OBJECTS IN ORDER AND COLLECTING PROPERLY!
                check=1
                #Get extra data here:
                if "CASENO." in split[i].replace(" ","").upper() and case_number=="":# Case Number & Document
                    case_number = split[i].upper().split("NO.")[1].replace(" ","").replace(":","")
                    #Get Document until can't
                    c=1
                    bypass=True
                    while bypass or ("ASSIGNEDTO" not in split[i+c].replace(" ","").upper() and "CUTOFF:" not in split[i+c].upper() and "PLACE:" not in split[i+c].upper() and "FILED:" not in split[i+c].upper() and "DATE:" not in split[i+c].upper() and c<20 and "PROPOUNDINGPARTY:" not in split[i+c].replace(" ","")):
                        val = min(max(0,len(split[i+c])-1),10)
                        if val>2:
                            if split[i+c][:val]==split[i+c][:val].upper() or bypass==False:
                                document = document + split[i+c]
                                bypass=False
                        c+=1
                elif "COUNTY" in split[i].replace(" ","") and county=="":# County
                    c=1
                    county = split[i]
                    #Get County until can't
                    while split[i+c].replace(" ","").replace("\n","")!="" and split[i+c]==split[i+c].upper() and c<10:
                        county=county+split[i+c]
                        c+=1
                #Ignore couunty if possible
                elif "Plaintiff" in split[i] and plaintiff=="" and county!="":
                    print("Founde")
                    c=1
                    #Get County until can't
                    while (split[i-c] not in county or len(split[i-c].replace(" ",""))<2) and c<20:
                        plaintiff = split[i-c] + plaintiff
                        c+=1
                elif "Defendant" in split[i] and defendant=="" and plaintiff!="":
                    print("def")
                    c=1
                    #Get County until can't
                    isEnd = str(re.sub(r'[^a-zA-Z]', '', str(split[i-c])))
                    while isEnd.replace(" ","").upper() not in ["V","VS"] and c<20:
                        defendant = split[i-c] + defendant
                        c+=1
                        isEnd = str(re.sub(r'[^a-zA-Z]', '', str(split[i-c])))
                elif "PROPOUNDINGPARTY:" in split[i].replace(" ","") or "REQUESTINGPARTY:" in split[i].replace(" ","") or "DEMANDINGPARTY:" in split[i].replace(" ","") and propounding_party=="":#Plaintiff and defendant
                    #Add until not propounding party
                    propounding_party=split[i].split(":")[-1]
                    c=1
                    while "SETNUMBER" not in split[i+c].replace(" ","") and "SETNO" not in split[i+c].replace(" ","")  and c<20:
                        if responding_party=="":# If adding to plaintiff
                            if "RESPONDINGPARTY:" in split[i+c].replace(" ",""):
                                responding_party = split[i+c].split(":")[-1]
                            else:
                                propounding_party = propounding_party + split[i+c]
                        else:
                            responding_party = responding_party + split[i+c]
                        c+=1
        except Exception as e:
            print(e)
            print("Error getting document details!")



        ##################################### 2. GET ACTUAL REQUESTS
        if not hard_stop:
            #print("Next")
            #print(split[i])
            if len(split[i].replace(" ",""))<50 and any(t in split[i][:min(len(split[i]),50)].replace(" ","").upper() for t in terms) and (split[i].replace(" ","")[-1] in [":","."] or split[i].replace(" ","")[-1].isdigit()):#       If request term used, must end in a certain character or a number, in case it is in text. Could check split length?
                #Add the custom key
                #print(split[i])
                key_matches =re.findall(r'\d+', split[i][min(10,len(split[i])):])
                if key_matches:
                    key = key_matches[0]
                    if start_at_one==False or (term_used==False and key!="1"):
                        keys.append(key)
                        start_at_one=False
                    adding=True
                    if req!="":
                        reqs.append(req.strip())
                    req=""
                    if not term_used:
                        reqs=[]
                    term_used=True

                
            elif "1. "==split[i][:3] and term_used==False:#                                    Restart requests
                reqs=[]
                adding=True
                req = split[i].split(".",1)[1]
            elif str(len(reqs)+check)+"." in split[i][:10] and term_used==False:#            If basic numbering used
                adding=True
                if req!="":
                    reqs.append(req.strip())
                req = split[i].split(".",1)[1]
            elif adding==True:#                                         Add the request text 
                if any(end in split[i].replace(" ","")[:10].upper() for end in ["DATED:","DATEDTHIS"]):
                    adding = False
                    hard_stop=True#Used to stop scanning for reqs
                elif len(split[i].replace(" ",""))>2 and not(split[i].replace(" ","")==split[i].replace(" ","").upper() and len(split[i].replace(" ",""))>30 and split[i].replace(" ","").strip()[-1] not in [".","?"]):
                    req = req+" "+re.sub(r"-[ 0-9 ]-","",split[i].strip().replace("///",""))


    #Add final request
    reqs.append(req.strip())

    ##################################### 3. GET TYPE OF DISCOVERY

    req_type=""
    search = data[:1000].replace(" ","")# Search 1st 1000 chars
    if "INTERROGATOR" in search:
        req_type="SPROG"
    elif "ADMISSION" in search:
        req_type="RFA"
    else:
        req_type="RFP"

    ##################################### 4. RETURN FINAL DATA

    #Remove 'county of' from county
    county = county.split("COUNTY OF ",1)[-1]
    #Clean propounding party and responding party
    #Remove plaintiff and defendant from the terms
    propounding_party = " ".join(propounding_party.split()[1:])
    responding_party = " ".join(responding_party.split()[1:])

    print("GOT HERE!")
    details = {"case_number":case_number.strip(),
                "document":document.strip(),
                "county":county.strip(),
                "plaintiff":plaintiff.strip(),
                "defendant":defendant.strip(),
                "propounding_party":propounding_party.strip(),
                "responding_party":responding_party.strip(),
                "date":""}

    return reqs,req_type,details,keys


### COMBINE ALL INTO GETREQUESTS
########################################################################################################
# Get the requests and filter
def getRequests(file):
    #If the user wants an empty frog file!
    if file=="BLANK FROG":
        reqs = FORM_VALUES
        reqs_type = "FROG"
        details = {"case_number":"",
                    "document":"",
                    "county":"",
                    "plaintiff":"",
                    "defendant":"",
                    "propounding_party":"",
                    "responding_party":"",
                    "date":""}
        keys=[]
        return reqs,reqs_type,details,keys
    #If a form interrogatory
    data = readPDF3(file)
    backup_data = readPDF(file)
    if "DISC-001" in data.replace(" ","")[:5000]:
        reqs = readFormThreaded(file)
        reqs_type = "FROG"
        details = {"case_number":"",
                    "document":"",
                    "county":"",
                    "plaintiff":"",
                    "defendant":"",
                    "propounding_party":"",
                    "responding_party":"",
                    "date":""}
        keys=[]
    else:
        reqs,reqs_type,details,keys = filterPDF(data)
        """
        breqs,breqs_type,bdetails,keys = filterPDF(backup_data)
        for i in range(len(reqs)):
            if reqs[i]=="":
                reqs[i]=breqs[i]
        """
    return reqs,reqs_type,details,keys


### UPLOAD THE DOCS
########################################################################################################
# Save requests, responses and details to a word DOCX
def updateDOC(reqs,resps,details,firm_details,file,name,numbers):
    templates = {
    "RFA":"assets/Response to RFA.docx",
    "RFP":"assets/Response to RFP.docx",
    "SPROG":"assets/Response to SROG.docx",
    "FROG":"assets/Response to FROG2.docx"
    }
    headings = {
    "RFA":"REQUEST FOR ADMISSION",
    "RFP":"DEMAND",
    "SPROG":"SPECIAL INTERROGATORY",
    "FROG":"FORM INTERROGATORY"
    }
    #Select the correct heading for the file
    if isinstance(file,str):
        heading=headings[file]
        #Loading template file
        filename=os.path.join(os.path.dirname(__file__),templates[file])
    else:#If a list of types (for export with client), set headings later
        filename=os.path.join(os.path.dirname(__file__),templates["RFA"])
    doc = Document(filename)
    #Normal Style
    if "FROG" in filename:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
    else:
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(12)
    #Heading style
    style = doc.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.bold=True
    font.underline=True

    # 1. ADD THE DETAILS
    #FOOTER
    for section in doc.sections:
        footer=section.footer
        footer.paragraphs[1].text  = footer.paragraphs[1].text.replace("DOCUMENTX", details["document"])
    #TABLE TEXT
    #First Column
    table = doc.tables[0].row_cells(0)
    for p in table[0].paragraphs:
        text = str(p.text)
        text=text.replace("PLAINTIFFX",details["plaintiff"])
        text=text.replace("DEFENDANTX",details["defendant"])
        p.text = text
    #Second column
    for p in table[1].paragraphs:
        text = str(p.text)
        if "CASENUMBERX" in text:
            #text=text.replace("CASE NUMBERX",details["case_number"])
            p.text=""
            run = p.add_run("Case No. "+str(details["case_number"]))
            run.bold=True
        elif "DOCUMENTX" in text:
            p.text=""
            run = p.add_run("PLAINTIFF'S RESPONSES TO "+str(details["document"]))
            run.bold=True

    #GENERAL TEXT
    for p in doc.paragraphs:
        text = str(p.text)
        if (x in text for x in ["PLAINTIFFX","DEFENDANTX","PROPOUNDINGX","RESPONDINGX","DOCUMENTX","DATEX","ATTORNEYSX","FIRM_NAMEX","ADDRESS_LINE_1X","ADDRESS_LINE_2X","TELEPHONEX","FACSIMILEX","EMAILX"]):
            #Normal Details
            text=text.replace("PLAINTIFFX",details["plaintiff"])
            text=text.replace("DEFENDANTX",details["defendant"])
            text=text.replace("PROPOUNDINGX",details["propounding_party"])
            text=text.replace("RESPONDINGX",details["responding_party"])
            text=text.replace("DOCUMENTX",details["document"])
            text=text.replace("DATEX",details["date"])
            #Firm Details
            text=text.replace("ATTORNEYSX",firm_details["attorneys"])
            text=text.replace("FIRM_NAMEX",firm_details["firm_name"])
            text=text.replace("ADDRESS_LINE_1X",firm_details["address_line_1"])
            text=text.replace("ADDRESS_LINE_2X",firm_details["address_line_2"])
            text=text.replace("TELEPHONEX",firm_details["telephone"])
            text=text.replace("FACSIMILEX",firm_details["facsimile"])
            text=text.replace("EMAILX",firm_details["email"])
            p.text = text
        if "NAME OF COUNTYX" in text:
            p.text=""
            run = p.add_run(text.replace("NAME OF COUNTYX",details["county"]))
            run.bold=True


    # 2. ADD THE REQUEST RESPONSES
    counter=0
    started=False
    for p in doc.paragraphs:
        if "STARTX" in p.text:
            started=True
            delete_paragraph(p)
        if started:
            for i in range(len(reqs)):#Add each paragraph with it's own correct heading
                if not isinstance(file,str):
                    heading = headings[file[counter]]
                #Add heading
                prev_p = insert_paragraph_after(prev_p,heading+" NO. "+str(numbers[counter])+":","Heading")
                #Add request
                prev_p = insert_paragraph_after(prev_p,"           "+reqs[counter],"Normal")
                #Add heading
                prev_p = insert_paragraph_after(prev_p,"RESPONSE TO "+heading+":","Heading")
                #Add response
                prev_p = insert_paragraph_after(prev_p,"           "+resps[counter],"Normal")
                #Next
                counter+=1
            doc.save(str(name)+".docx")
            return
        prev_p = p






### READ CLIENT FEEDBACK
########################################################################################################
def read_client_feedback(filename):
    #Read client feedback from a DOCX file
    #Given this filename, open, get data into sections. Send to main program to update requests
    #Each one needs key,resp,type
    #EXAMPLE FORMAT: feedback=[{"key":"1","type":"RFA","response":"CLIENT FEEDBACK: incredible"}]
    feedback=[]
    doc = Document(filename)
    #Get each key from the request
    #Then find the response and add to a new dict
    current_key=""
    current_type=""
    current_response=""
    for p in doc.paragraphs:
        #RFA
        if "REQUEST FOR ADMISSION NO. " in p.text:
            if current_response!="":
                feedback.append({"key":current_key,"type":current_type,"response":"CLIENT FEEDBACK: "+str(current_response)})
            current_key = ".".join(re.findall(r'\d+', p.text))
            current_type = "RFA"
        elif "RESPONSE TO REQUEST FOR ADMISSION:" in p.text or "RESPONSE TO DEMAND:" in p.text or "RESPONSE TO SPECIAL INTERROGATORY:" in p.text or "RESPONSE TO FORM INTERROGATORY:" in p.text:
            current_response = ""
        #RFP
        elif "DEMAND NO. " in p.text:
            if current_response!="":
                feedback.append({"key":current_key,"type":current_type,"response":"CLIENT FEEDBACK: "+str(current_response)})
            current_key = ".".join(re.findall(r'\d+', p.text))
            current_type = "RFP"
        #SPROG
        elif "SPECIAL INTERROGATORY NO. " in p.text:
            if current_response!="":
                feedback.append({"key":current_key,"type":current_type,"response":"CLIENT FEEDBACK: "+str(current_response)})
            current_key = ".".join(re.findall(r'\d+', p.text))
            current_type = "SPROG"
        #FROG
        elif "FORM INTERROGATORY NO. " in p.text or "Dated:" in p.text[:8]:
            if current_response!="":
                feedback.append({"key":current_key,"type":current_type,"response":"CLIENT FEEDBACK: "+str(current_response)})
            current_key = ".".join(re.findall(r'\d+', p.text))
            current_type = "FROG"
        else:
            current_response = current_response + p.text#Add current paragraph
    return feedback




##SHOULD ADD THREADS FOR BOX DETECT!