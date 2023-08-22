import PyPDF2 as pdf
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt
from pdfminer.high_level import extract_text
import fitz
#Types of discovery requests
#SPECIAL INTERROGATORY
#REQUEST FOR ADMISSIONS
#REQUEST FOR PRODUCTION OF DOCUMENTS / DEMAND FOR PRODUCTION OF DOCUMENTS

#Create a word doc using the PDF and templates to offer responses to all requests


#Reads a pdf
def readPDF(file):
    reader = pdf.PdfReader(file)# Create reader object
    text=[]# Get each page
    for i in reader.pages:
        text.append(i.extract_text())
    text = "".join(list(text))
    return text

def readPDF2(file):
    text = extract_text(file)
    return text

def readPDF3(file):
    doc = fitz.open(file)
    text = ""
    for page in doc:
        text+=page.get_text()
    return text







#Gets only the needed requests from the PDF
def filterPDF(data):
    #NEED TO FIND A SEQUENCE OF REQUESTS!
    print("Document Length: "+str(len(data))+" pages")
    #data = list(data)
    #data = " ".join(data)
    
    #Search Terms
    terms=["REQUESTNO.","INTERROGATORYNO.","ANDNO.","ENTSNO.","ONSNO.","TIONNO."]
    reqs=[]
    req=""
    adding=False
    split = data.split("\n")

    ####################### MAIN LOOP

    # Document vars
    case_number=""
    county=""
    plaintiff=""
    defendant=""
    document=""

    for i in range(len(split)):
        check=2
        if adding==False:#THIS KEEPS THE OBJECTS IN ORDER AND COLLECTING PROPERLY!
            check=1
            #Get extra data here:
            if "CASENO." in split[i].replace(" ","").upper() and case_number=="":# Case Number & Document
                case_number = split[i].upper().split("NO.")[1].replace(" ","").replace(":","")
                #Get Document until can't
                c=1
                bypass=True
                while bypass or ("ASSIGNEDTO" not in split[i+c].replace(" ","").upper() and "CUTOFF:" not in split[i+c].upper() and "PLACE:" not in split[i+c].upper() and "FILED:" not in split[i+c].upper() and "DATE:" not in split[i+c].upper() and c<20 and "PROPOUNDINGPARTY:" not in split[i+c].replace(" ","")):
                    val = min(max(0,len(split[i+c])-1),10)
                    if val>2:
                        if split[i+c][:val]==split[i+c][:val].upper() or bypass==False:
                            document = document + split[i+c]
                            bypass=False
                    c+=1

            elif "COUNTY" in split[i].replace(" ","") and county=="":# County
                county = split[i].split("COUNTY OF",1)[-1]
                c=1
                #Get County until can't
                while split[i+c].replace(" ","").replace("\n","")!="" and split[i+c]==split[i+c].upper() and c<10:
                    county=county+split[i+c]
                    c+=1

            elif "PROPOUNDINGPARTY:" in split[i].replace(" ","") and plaintiff=="":#Plaintiff and defendant
                #Add until not propounding party
                plaintiff=split[i].split(":")[-1]
                c=1
                while "SETNUMBER" not in split[i+c].replace(" ","") and "SETNO" not in split[i+c].replace(" ","")  and c<20:
                    if defendant=="":# If adding to plaintiff
                        if "RESPONDINGPARTY:" in split[i+c].replace(" ",""):
                            defendant = split[i+c].split(":")[-1]
                        else:
                            plaintiff = plaintiff + split[i+c]
                    else:
                        defendant = defendant + split[i+c]
                    c+=1

        #Clean Plaintiff and defendant
        plaintiff = plaintiff.upper().replace("DEFENDANT","").replace("(S)","")
        done = False
        for i2 in range(len(plaintiff)):
            if done==False:
                if plaintiff[i2].isalpha():
                    plaintiff = plaintiff[i2:]
                    done = True

        defendant = defendant.upper().replace("PLAINTIFF","").replace("(S)","")
        done = False
        for i2 in range(len(defendant)):
            if done==False:
                if defendant[i2].isalpha():
                    defendant = defendant[i2:]
                    done = True

        #SWAP THEM!!! TEMP
        temp = plaintiff
        plaintiff = defendant
        defendant = temp
            
        # Get main data here
        if any(t in split[i].replace(" ","") for t in terms):# If request term used
            adding=True
            if req!="":
                reqs.append(req.strip())
            req=""

        elif str(len(reqs)+check)+". " in split[i][:10]:# If basic numbering used
            adding=True
            if req!="":
                reqs.append(req.strip())
            req = split[i].split(".",1)[1]
        elif adding==True:# Add the request text 
            if "DATED:" in split[i].replace(" ","")[:10].upper():
                adding = False
            elif len(split[i].replace(" ",""))>2 and not(split[i].replace(" ","")==split[i].replace(" ","").upper() and  "." not in split[i]):
                req = req+" "+split[i].strip()


    reqs.append(req.strip())
    #######################

    print("Returned Requests: "+str(len(reqs))+"\n")

    """
    c=1
    for i in reqs:
        print("REQUEST "+str(c))
        print(i)
        print("")
        c+=1
    """

    #Get type of discovery
    req_type=""
    search = data[:1000].replace(" ","")# Search 1st 1000 chars
    if "INTERROGATORIES" in search:
        req_type="SPROG"
    elif "ADMISSIONS" in search:
        req_type="RFA"
    else:
        req_type="RFP"
    
    """
    print("CASE: "+str(case_number))
    print("DOCUMENT: "+str(document))
    print("COUNTY: "+str(county))
    print("Plaintiff: "+str(plaintiff))
    print("Defendant: "+str(defendant))
    """
    details = {"case_number":case_number,
                "document":document,
                "county":county,
                "plaintiff":plaintiff,
                "defendant":defendant}

    return reqs,req_type,details



def getRequests(file):
    data = readPDF3(file)
    backup_data = readPDF(file)

    reqs,reqs_type,details = filterPDF(data)
    breqs,breqs_type,bdetails = filterPDF(backup_data)

    for i in range(len(reqs)):
        if reqs[i]=="":
            reqs[i]=breqs[i]
    return reqs,reqs_type,details





def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style

    return new_para

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None



def updateDOC(reqs,resps,details,file,name):
    templates = {
    "RFA":"assets/Response to RFA.docx",
    "RFP":"assets/Response to RFP.docx",
    "SPROG":"assets/Response to SROG.docx"
    }
    #Loading template file
    file=templates[file]
    doc = Document(file)
    #Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)



    # 1. ADD THE DETAILS
    # We have details{}

    #Footer
    for section in doc.sections:
        footer=section.footer
        footer.paragraphs[1].text  = footer.paragraphs[1].text.replace("DOCUMENTX", details["document"])

    #Initial Table
    table = doc.tables[0].row_cells(0)
    for p in table[0].paragraphs:
        text = str(p.text)
        text=text.replace("PLAINTIFFX",details["plaintiff"])
        text=text.replace("DEFENDANTX",details["defendant"])
        p.text = text

    for p in table[1].paragraphs:
        text = str(p.text)
        if "CASE NUMBERX" in text:
            #text=text.replace("CASE NUMBERX",details["case_number"])
            p.text=""
            run = p.add_run("Case No. "+str(details["case_number"]))
            run.bold=True
        elif "DOCUMENTX" in text:
            p.text=""
            run = p.add_run("PLAINTIFF’S RESPONSES TO "+str(details["document"]))
            run.bold=True
        #text=text.replace("DOCUMENTX",details["document"])
        #p.text = text 



    #General text
    for p in doc.paragraphs:
        text = str(p.text)
        if "PLAINTIFFX" in text or "DEFENDANTX" in text or "DOCUMENTX" in text:
            text=text.replace("PLAINTIFFX",details["plaintiff"])
            text=text.replace("DEFENDANTX",details["defendant"])
            text=text.replace("DOCUMENTX",details["document"])
            p.text = text
        if "NAME OF COUNTYX" in text:
            p.text=""
            run = p.add_run(text.replace("NAME OF COUNTYX",details["county"]))
            run.bold=True
            

    #Find exact doc for each to go in!

    # 2. ADD THE REQUEST RESPONSES
    counter=0
    for p in doc.paragraphs:
        if "NO. " in p.text:
            if counter<len(reqs):
                if "RESPONSE" not in p.text:
                    insert_paragraph_after(p,"           "+reqs[counter],"Normal")
                    
                else:
                    insert_paragraph_after(p,"           "+resps[counter],"Normal")
                    counter+=1
            else:
                #Destroy para
                if "RESPONSE" not in p.text:
                    delete_paragraph(prev_p)
                delete_paragraph(p)
        prev_p = p
            #Add the next item! If not response
            #Remove if done


    doc.save(str(name)+".docx")

